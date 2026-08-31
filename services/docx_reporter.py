"""Word report for a completed source-versus-target validation run."""

from pathlib import Path
import logging

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, RGBColor

from services.dashboard_inventory_service import (
    _count_inventory_for_execution,
    _list_kpis_for_execution,
    _list_visuals_for_execution,
)


LIGHT_RED = "FCE4D6"
logger = logging.getLogger(__name__)


def _shade(cell, color):
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    properties.append(shading)


def _table(document, headers, rows):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = str(header)
        _shade(cell, LIGHT_RED)
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(156, 0, 6)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = "" if value is None else str(value)
    return table


def _status_brief(items, label):
    total = len(items)
    matches = sum(1 for item in items if item.get("status") == "Match")
    differences = total - matches
    return f"{label}: {matches} matching and {differences} mismatching or missing out of {total} compared item(s)."


def _browser_metric_rows(source_metrics, target_metrics):
    rows = []
    for label, key in [
        ("Page name", "page_name"),
        ("Page load (seconds)", "page_load_seconds"),
        ("Render (seconds)", "dashboard_render_seconds"),
        ("Filter dashboard render (seconds)", "filter_dashboard_render_seconds"),
        ("Screenshot (seconds)", "screenshot_seconds"),
        ("Gemini extraction (seconds)", "gemini_extraction_seconds"),
        ("Total execution (seconds)", "total_execution_seconds"),
        ("HTTP requests", "total_requests"),
        ("Failed requests", "failed_requests"),
        ("Console messages", "console_messages"),
        ("Page errors", "page_errors"),
    ]:
        source_value = source_metrics.get(key)
        target_value = target_metrics.get(key)
        if source_value is None and target_value is None:
            continue
        rows.append((label, source_value, target_value))
    return rows


def generate_validation_document(
    run_id,
    executions,
    comparison,
    output_directory,
    page_comparisons=None,
    executions_by_dashboard=None,
):
    output_directory = Path(output_directory)
    output_directory.mkdir(parents=True, exist_ok=True)
    document = Document()
    
    for style_name in ("Heading 1", "Heading 2"):
        document.styles[style_name].font.color.rgb = RGBColor(156, 0, 6)
        
    section = document.sections[0]
    section.top_margin = section.bottom_margin = Inches(0.7)
    
    title = document.add_heading("Power BI Dashboard Validation Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(156, 0, 6)
        
    document.add_paragraph(f"Validation run: {run_id}")

    document.add_heading("Validation Status", level=1)
    document.add_paragraph(comparison.get("status", "not_compared").replace("_", " ").title())
    if comparison.get("reason"):
        document.add_paragraph(comparison["reason"])

    if page_comparisons:
        document.add_heading("Multi-Page Comparison Summary", level=1)
        _table(
            document,
            ["Page Name", "Status", "Overall %", "Filter %", "KPI %", "Visual %"],
            [
                (
                    item.get("page_name"),
                    item.get("status"),
                    (item.get("summary") or {}).get("overall_match_percentage"),
                    (item.get("summary") or {}).get("filter_match_percentage"),
                    (item.get("summary") or {}).get("kpi_match_percentage"),
                    (item.get("summary") or {}).get("visual_match_percentage"),
                )
                for item in page_comparisons
            ],
        )

    document.add_heading("Dashboard Metadata", level=1)
    _table(document, ["Field", "Source", "Target"], [
        ("Name", executions[0]["dashboard"].get("name"), executions[1]["dashboard"].get("name")),
        ("URL", executions[0]["dashboard"].get("url"), executions[1]["dashboard"].get("url")),
        ("Page", executions[0]["dashboard"].get("page_name"), executions[1]["dashboard"].get("page_name")),
        ("Browser visual extraction", executions[0]["visual_data"].get("status"), executions[1]["visual_data"].get("status")),
    ])

    if comparison.get("status") == "success":
        document.add_heading("Validation Match Summary", level=1)
        summary = comparison["summary"]
        _table(document, ["Area", "Match %"], [
            ("Filters", summary.get("filter_match_percentage")), ("KPIs", summary.get("kpi_match_percentage")),
            ("Visuals", summary.get("visual_match_percentage")), ("Overall", summary.get("overall_match_percentage")),
        ])

    source_metrics = executions[0]["metrics"] or {}
    target_metrics = executions[1]["metrics"] or {}
    document.add_heading("Browser Metrics (First Compared Page)", level=1)
    metric_rows = _browser_metric_rows(source_metrics, target_metrics)
    if metric_rows:
        _table(document, ["Metric", "Source", "Target"], metric_rows)
    else:
        document.add_paragraph("No browser metrics were captured for the first compared page.")

    if executions_by_dashboard:
        document.add_heading("Page Inventory (All Pages)", level=1)
        inventory_rows = []
        for dashboard_executions in executions_by_dashboard:
            for execution in dashboard_executions:
                dashboard = execution.get("dashboard") or {}
                inventory = _count_inventory_for_execution(execution)
                inventory_rows.append((
                    dashboard.get("name"),
                    dashboard.get("page_name") or inventory.get("page_name") or "Default",
                    inventory.get("filter_count", 0), inventory.get("kpi_count", 0),
                    inventory.get("table_count", 0), inventory.get("matrix_count", 0),
                    inventory.get("chart_count", 0), inventory.get("total_visuals", 0),
                ))
        _table(document, ["Dashboard", "Page", "Filters", "KPIs", "Tables", "Matrices", "Charts", "Total Visuals"], inventory_rows)

    if comparison.get("status") == "success":
        document.add_heading("KPI Comparison", level=1)
        kpis = comparison.get("kpis", [])
        if kpis:
            _table(document, ["KPI", "Source", "Target", "Status"], [
                (item.get("kpi", item.get("name")), item.get("source", item.get("source_value")), item.get("target", item.get("target_value")), item.get("status"))
                for item in kpis
            ])
            document.add_paragraph(_status_brief(kpis, "KPI result"))
        else:
            document.add_paragraph("No KPI cards were compared.")

        document.add_heading("Visual Comparison", level=1)
        visuals = comparison.get("visuals", [])
        if visuals:
            _table(document, ["Visual", "Source Data Points", "Target Data Points", "Status"], [
                (item.get("visual", item.get("title")), str(item.get("source", "")), str(item.get("target", "")), item.get("status"))
                for item in visuals
            ])
            document.add_paragraph(_status_brief(visuals, "Visual result"))
        else:
            document.add_paragraph("No visuals were compared.")

    path = output_directory / f"{run_id}_dashboard_validation.docx"
    document.save(path)
    logger.info("Word validation report generated | %s", path)
    return path